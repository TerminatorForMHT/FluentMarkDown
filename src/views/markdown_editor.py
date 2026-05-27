import os
import markdown

from PyQt5.QtWidgets import (
    QTextEdit, QVBoxLayout, QFrame, QWidget, QStatusBar, QFileDialog
)
from PyQt5.QtCore import Qt,  QPoint
from PyQt5.QtGui import QPainterPath, QRegion, QColor

from qfluentwidgets import (
    FluentIcon,
    CommandBar,
    TransparentPushButton,
    CardWidget,
    ComboBox
    )
import os
import markdown

from PyQt5.QtWidgets import (
    QTextEdit, QVBoxLayout, QFrame, QWidget, QStatusBar, QFileDialog
)
from PyQt5.QtCore import Qt, QPoint, QTimer, QCryptographicHash
from PyQt5.QtGui import QPainterPath, QRegion, QColor

from qfluentwidgets import (
    FluentIcon,
    CommandBar,
    TransparentPushButton,
    CardWidget,
    ComboBox,
    BodyLabel,
    SingleDirectionScrollArea,
    isDarkTheme,
    Dialog,
    MessageBox,
    DropDownPushButton,
    RoundMenu,
    Action
)
from qframelesswindow.webengine import FramelessWebEngineView

import json
import os
from src.models.themes import PreviewThemes

# 导出功能所需的库
try:
    from fpdf import FPDF
    from docx import Document
    HAS_EXPORT_LIBS = True
except ImportError:
    HAS_EXPORT_LIBS = False


class MarkdownWidget(QFrame):
    """Markdown 编辑和预览界面"""

    PREVIEW_RADIUS = 8
    PREVIEW_UPDATE_DELAY = 300

    def __init__(self, parent=None):
        super().__init__(parent=parent)
        self.setObjectName("markdownInterface")
        self.is_fullscreen = False
        self.is_editor_fullscreen = False
        self.has_file = False
        self.is_modified = False
        self.current_file_path = None
        self.preview_theme = "light"
        self.font_size = 16
        self._preview_updating = False
        self._preview_dirty = False
        self._cached_html_template = None
        self._last_md5 = None
        self._command_bar_buttons = {}
        self._auto_save_timer = QTimer(self)
        self._auto_save_timer.timeout.connect(self._auto_save)
        self._auto_save_interval = 30000
        self._history_file_path = os.path.join(os.path.expanduser("~"), ".fluentmarkdown_history.json")
        self._recent_files = self._load_recent_files()
        self._history_menu = None

        # 主布局
        self.vBoxLayout = QVBoxLayout(self)
        self.vBoxLayout.setContentsMargins(0, 0, 0, 0)
        self.vBoxLayout.setSpacing(0)

        # 命令栏
        self.command_bar = CommandBar(self)
        self.setup_command_bar()

        # 容器
        self.card_container = QWidget(self)
        self.card_container_layout = QVBoxLayout(self.card_container)
        self.card_container_layout.setContentsMargins(5, 5, 5, 5)
        self.card_container_layout.setSpacing(0)

        # 编辑/预览卡片（外层仍然用 CardWidget 没问题）
        self.editor_card = CardWidget(self.card_container)
        self.editor_card.setStyleSheet("background-color: transparent;")
        self.editor_layout = QVBoxLayout(self.editor_card)
        self.editor_layout.setContentsMargins(1, 1, 1, 1)
        self.editor_layout.setSpacing(0)

        # splitter
        from PyQt5.QtWidgets import QSplitter
        self.splitter = QSplitter(Qt.Horizontal, self.editor_card)
        self.splitter.setStyleSheet('''
            QSplitter { background-color: transparent; }
            QSplitter::handle { width: 0px; background-color: transparent; }
        ''')
        self.splitter.setHandleWidth(0)
        self.splitter.setEnabled(False)
        self.editor_layout.addWidget(self.splitter, 1)
        self.card_container_layout.addWidget(self.editor_card, 1)

        # 左侧编辑器
        self.editor = QTextEdit()
        self.editor.setPlaceholderText("Write Markdown here...")
        self.editor.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.editor.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        # 安装事件过滤器，捕获编辑器的鼠标滚轮事件
        self.editor.installEventFilter(self)

        self.scroll_area = SingleDirectionScrollArea(orient=Qt.Vertical, parent=self.editor_card)
        self.scroll_area.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        self.scroll_area.setWidget(self.editor)
        self.scroll_area.setWidgetResizable(True)

        # 右侧预览容器：关键修复点 —— 不用 CardWidget，避免它自己画白底/阴影挡住 Mica
        self.preview_container = QFrame(self.editor_card)
        self.preview_container.setObjectName("previewContainer")
        self.preview_container.setAttribute(Qt.WA_TranslucentBackground, True)
        self.preview_container.setStyleSheet("""
            QFrame#previewContainer{
                background: transparent;
                border: none;
            }
        """)
        self.preview_layout = QVBoxLayout(self.preview_container)
        self.preview_layout.setContentsMargins(0, 0, 0, 0)
        self.preview_layout.setSpacing(0)

        # WebEngine 预览
        self.preview = FramelessWebEngineView(self.preview_container)
        self.preview.setAttribute(Qt.WA_TranslucentBackground, True)
        self.preview.setStyleSheet("background: transparent; border: none;")
        # 关键：让 WebEngine 页面背景真正透明（否则默认白）
        try:
            self.preview.page().setBackgroundColor(QColor(0, 0, 0, 0))
        except Exception:
            pass

        self.preview_layout.addWidget(self.preview)

        # 加入 splitter
        self.splitter.addWidget(self.scroll_area)
        self.splitter.addWidget(self.preview_container)
        self.splitter.setSizes([500, 500])
        self.splitter.setStretchFactor(0, 1)
        self.splitter.setStretchFactor(1, 1)

        # 状态栏
        self.status_bar = QStatusBar(self)
        from qfluentwidgets import isDarkTheme
        is_dark = isDarkTheme()
        if is_dark:
            self.status_bar.setStyleSheet("""
                QStatusBar {
                    background-color: transparent;
                    border: none;
                    padding: 2px 8px;
                }
                QStatusBar::item {
                    border: none;
                }
            """)
        else:
            self.status_bar.setStyleSheet("""
                QStatusBar {
                    background-color: transparent;
                    border: none;
                    padding: 2px 8px;
                }
                QStatusBar::item {
                    border: none;
                }
            """)
        self.setup_status_bar()

        # 主布局装载
        self.vBoxLayout.addWidget(self.command_bar)
        self.vBoxLayout.addWidget(self.card_container, 1)
        self.vBoxLayout.addWidget(self.status_bar)

        # 防抖定时器
        self._preview_timer = QTimer(self)
        self._preview_timer.setSingleShot(True)
        self._preview_timer.timeout.connect(self._do_preview_update)

        # 信号（使用防抖）
        self.editor.textChanged.connect(self._on_text_changed)
        self.editor.selectionChanged.connect(self.update_status_bar)

        # 初始化
        self.update_editor_style()
        self.update_status_bar()

        # 初始状态：显示欢迎页
        QTimer.singleShot(0, self._show_welcome_page)
        QTimer.singleShot(10, self._update_command_bar_enabled)

    # ---------------- 圆角裁剪（只裁右上/右下） ----------------
    def _show_welcome_page(self):
        self.has_file = False
        self.editor.clear()
        self.editor.hide()
        from qfluentwidgets import isDarkTheme
        is_dark = isDarkTheme()
        welcome_html = self._generate_welcome_html(is_dark)
        self.preview.setHtml(welcome_html)
        self.splitter.setSizes([0, self.splitter.width()])

    def _generate_welcome_html(self, is_dark):
        if is_dark:
            body_color = "rgba(255, 255, 255, 0.9)"
            h1_color = "rgba(255, 255, 255, 0.85)"
            p_color = "rgba(255, 255, 255, 0.6)"
            shortcut_bg = "rgba(255, 255, 255, 0.1)"
            shortcut_key_bg = "rgba(255, 255, 255, 0.15)"
            shortcut_key_color = "rgba(255, 255, 255, 0.9)"
        else:
            body_color = "rgba(0, 0, 0, 0.9)"
            h1_color = "rgba(0, 0, 0, 0.85)"
            p_color = "rgba(0, 0, 0, 0.6)"
            shortcut_bg = "rgba(0, 0, 0, 0.06)"
            shortcut_key_bg = "rgba(0, 0, 0, 0.1)"
            shortcut_key_color = "rgba(0, 0, 0, 0.8)"
        return f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            background: transparent;
            color: {body_color};
        }}
        .container {{
            text-align: center;
            padding: 40px;
        }}
        .icon {{
            font-size: 72px;
            margin-bottom: 24px;
        }}
        h1 {{
            font-size: 36px;
            font-weight: 600;
            margin-bottom: 16px;
            letter-spacing: 1px;
            color: {h1_color};
        }}
        p {{
            font-size: 18px;
            margin-bottom: 32px;
            line-height: 1.6;
            color: {p_color};
        }}
        .shortcuts {{
            display: flex;
            gap: 24px;
            justify-content: center;
            flex-wrap: wrap;
        }}
        .shortcut {{
            background: {shortcut_bg};
            padding: 12px 20px;
            border-radius: 8px;
        }}
        .shortcut-key {{
            background: {shortcut_key_bg};
            padding: 4px 10px;
            border-radius: 4px;
            font-weight: 600;
            margin-right: 8px;
            color: {shortcut_key_color};
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="icon">📝</div>
        <h1>FluentMarkDown</h1>
        <p>新建或打开 Markdown 文件开始编辑</p>
        <div class="shortcuts">
            <div class="shortcut">
                <span class="shortcut-key">Ctrl+N</span>新建文件
            </div>
            <div class="shortcut">
                <span class="shortcut-key">Ctrl+O</span>打开文件
            </div>
        </div>
    </div>
</body>
</html>
'''

    def _load_recent_files(self):
        try:
            if os.path.exists(self._history_file_path):
                with open(self._history_file_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception:
            pass
        return []

    def _save_recent_files(self):
        try:
            with open(self._history_file_path, 'w', encoding='utf-8') as f:
                json.dump(self._recent_files, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _add_to_recent_files(self, file_path):
        if not file_path:
            return
        if file_path in self._recent_files:
            self._recent_files.remove(file_path)
        self._recent_files.insert(0, file_path)
        if len(self._recent_files) > 10:
            self._recent_files = self._recent_files[:10]
        self._save_recent_files()
        self._update_history_menu()

    def _update_history_menu(self):
        if self._history_menu is None:
            return
        self._history_menu.clear()
        if not self._recent_files:
            action = Action("无历史文件", self._history_menu)
            action.setEnabled(False)
            self._history_menu.addAction(action)
        else:
            for file_path in self._recent_files:
                file_name = os.path.basename(file_path)
                action = Action(file_name, self._history_menu)
                action.setToolTip(file_path)
                action.triggered.connect(lambda checked, fp=file_path: self.open_file(fp))
                self._history_menu.addAction(action)
            self._history_menu.addSeparator()
            clear_action = Action("清除历史", self._history_menu)
            clear_action.triggered.connect(self._clear_recent_files)
            self._history_menu.addAction(clear_action)

    def _clear_recent_files(self):
        self._recent_files = []
        self._save_recent_files()
        self._update_history_menu()

    def _show_info_dialog(self, title, content):
        w = MessageBox(title, content, self)
        w.exec()

    def _show_yes_no_dialog(self, title, content):
        w = MessageBox(title, content, self)
        return w.exec()

    def _auto_save(self):
        if self.has_file and self.is_modified and self.current_file_path:
            try:
                with open(self.current_file_path, 'w', encoding='utf-8') as f:
                    f.write(self.editor.toPlainText())
                self.is_modified = False
            except Exception as e:
                pass

    def _start_auto_save_timer(self):
        self._auto_save_timer.start(self._auto_save_interval)

    def _stop_auto_save_timer(self):
        self._auto_save_timer.stop()

    def check_save_on_close(self):
        if self.has_file and self.is_modified:
            ret = self._show_yes_no_dialog(
                "文件已修改",
                f"是否保存对 {self.current_file_path or 'untitled.md'} 的更改？"
            )
            if ret:
                self.save_file()
                return True
            else:
                return True
        return False

    def _hide_welcome_page(self):
        self._preview_updating = False
        self.editor.show()
        self.splitter.setEnabled(True)
        w = self.splitter.width()
        self.splitter.setSizes([w // 2, w // 2])

    def _update_command_bar_enabled(self):
        enabled = self.has_file
        for name, btn in self._command_bar_buttons.items():
            if btn and name not in ('new', 'open'):
                btn.setEnabled(enabled)

    def _updatePreviewRoundMask(self):
        if not hasattr(self, "preview_container"):
            return

        if self.is_fullscreen:
            self.preview_container.clearMask()
            return

        r = self.PREVIEW_RADIUS
        rect = self.preview_container.rect()
        if rect.isNull():
            return

        path = QPainterPath()
        path.moveTo(rect.topLeft())
        path.lineTo(rect.topRight() - QPoint(r, 0))
        path.quadTo(rect.topRight(), rect.topRight() + QPoint(0, r))
        path.lineTo(rect.bottomRight() - QPoint(0, r))
        path.quadTo(rect.bottomRight(), rect.bottomRight() - QPoint(r, 0))
        path.lineTo(rect.bottomLeft())
        path.closeSubpath()

        region = QRegion(path.toFillPolygon().toPolygon())
        self.preview_container.setMask(region)

    def resizeEvent(self, e):
        super().resizeEvent(e)
        self._updatePreviewRoundMask()
    
    def wheelEvent(self, event):
        """处理鼠标滚轮事件，实现Ctrl+滚轮调节字体大小"""
        from PyQt5.QtCore import Qt
        if event.modifiers() & Qt.ControlModifier:
            # Ctrl+滚轮调节字体大小
            delta = event.angleDelta().y()
            if delta > 0:
                # 放大字体
                self.font_size = min(self.font_size + 2, 32)  # 最大32px
            else:
                # 缩小字体
                self.font_size = max(self.font_size - 2, 8)   # 最小8px
            
            # 更新编辑器和预览的字体大小
            self.update_editor_style()
            self.update_preview()
        else:
            # 正常滚轮事件
            super().wheelEvent(event)

    # ---------------- 命令栏 ----------------
    def setup_command_bar(self):
        new_button = TransparentPushButton(FluentIcon.ADD, "新建")
        new_button.clicked.connect(self.new_file)
        self.command_bar.addWidget(new_button)
        self._command_bar_buttons['new'] = new_button

        open_button = TransparentPushButton(FluentIcon.FOLDER, "打开")
        open_button.clicked.connect(self.open_file_dialog)
        self.command_bar.addWidget(open_button)
        self._command_bar_buttons['open'] = open_button

        self._history_menu = RoundMenu(parent=self)
        self._update_history_menu()
        recent_button = DropDownPushButton(FluentIcon.HISTORY, "最近", self)
        recent_button.setMenu(self._history_menu)
        self.command_bar.addWidget(recent_button)

        save_button = TransparentPushButton(FluentIcon.SAVE, "保存")
        save_button.clicked.connect(self.save_file_dialog)
        self.command_bar.addWidget(save_button)
        self._command_bar_buttons['save'] = save_button

        self.command_bar.addSeparator()

        copy_button = TransparentPushButton(FluentIcon.COPY, "复制")
        copy_button.clicked.connect(self.copy)
        self.command_bar.addWidget(copy_button)
        self._command_bar_buttons['copy'] = copy_button

        paste_button = TransparentPushButton(FluentIcon.PASTE, "粘贴")
        paste_button.clicked.connect(self.paste)
        self.command_bar.addWidget(paste_button)
        self._command_bar_buttons['paste'] = paste_button

        self.command_bar.addSeparator()

        fullscreen_button = TransparentPushButton(FluentIcon.ZOOM_IN, "全屏阅读")
        fullscreen_button.clicked.connect(self.toggle_fullscreen)
        self.command_bar.addWidget(fullscreen_button)
        self._command_bar_buttons['fullscreen'] = fullscreen_button

        fullscreen_edit_button = TransparentPushButton(FluentIcon.EDIT, "全屏编辑")
        fullscreen_edit_button.clicked.connect(self.toggle_editor_fullscreen)
        self.command_bar.addWidget(fullscreen_edit_button)
        self._command_bar_buttons['fullscreen_edit'] = fullscreen_edit_button

        self.command_bar.addSeparator()

        self.setup_theme_menu()

        self.command_bar.addSeparator()

        image_button = TransparentPushButton(FluentIcon.PHOTO, "插入图片")
        image_button.clicked.connect(self.insert_image)
        self.command_bar.addWidget(image_button)

        self.command_bar.addSeparator()

        # 字体大小调节
        zoom_in_button = TransparentPushButton(FluentIcon.ZOOM_IN, "放大")
        zoom_in_button.clicked.connect(self.zoom_in)
        self.command_bar.addWidget(zoom_in_button)

        zoom_out_button = TransparentPushButton(FluentIcon.ZOOM_OUT, "缩小")
        zoom_out_button.clicked.connect(self.zoom_out)
        self.command_bar.addWidget(zoom_out_button)

        zoom_reset_button = TransparentPushButton(FluentIcon.HOME, "重置")
        zoom_reset_button.clicked.connect(self.zoom_reset)
        self.command_bar.addWidget(zoom_reset_button)

        self.command_bar.addSeparator()

        export_button = TransparentPushButton(FluentIcon.SHARE, "导出")
        export_button.clicked.connect(self.export_file)
        self.command_bar.addWidget(export_button)
        self._command_bar_buttons['export'] = export_button

    def setup_theme_menu(self):
        self.theme_label_cmd = BodyLabel("预览主题:")
        self.theme_combo = ComboBox()

        themes = PreviewThemes.get_available_themes()
        theme_names = []
        current_index = 0
        for i, theme in enumerate(themes):
            theme_info = PreviewThemes.get_theme_styles(theme)
            theme_names.append(theme_info["name"])
            if theme == self.preview_theme:
                current_index = i

        self.theme_combo.addItems(theme_names)
        self.theme_combo.setCurrentIndex(current_index)
        self.theme_combo.currentIndexChanged.connect(self.on_theme_changed)

        self.command_bar.addWidget(self.theme_label_cmd)
        self.command_bar.addWidget(self.theme_combo)

    def on_theme_changed(self, index):
        themes = PreviewThemes.get_available_themes()
        if 0 <= index < len(themes):
            self.set_preview_theme(themes[index])

    def set_preview_theme(self, theme_name):
        self.preview_theme = theme_name
        self._last_md5 = None
        self._cached_html_template = None
        self.update_preview()
        self.update_status_bar()

        themes = PreviewThemes.get_available_themes()
        if theme_name in themes:
            self.theme_combo.blockSignals(True)
            self.theme_combo.setCurrentIndex(themes.index(theme_name))
            self.theme_combo.blockSignals(False)

    # ---------------- 状态栏 ----------------
    def setup_status_bar(self):
        from qfluentwidgets import isDarkTheme
        is_dark = isDarkTheme()
        
        # 创建标签
        self.char_count_label = BodyLabel("字符: 0", self)
        self.selection_label = BodyLabel("选中: 0", self)
        self.encoding_label = BodyLabel("编码: UTF-8", self)

        theme_info = PreviewThemes.get_theme_styles(self.preview_theme)
        self.theme_label = BodyLabel(f"预览主题: {theme_info['name']}", self)
        
        # 设置标签样式，适配主题颜色并调整间距
        text_color = "#ffffff" if is_dark else "#333333"
        label_style = f"color: {text_color}; padding: 0 8px;"
        self.char_count_label.setStyleSheet(label_style)
        self.selection_label.setStyleSheet(label_style)
        self.theme_label.setStyleSheet(label_style)
        self.encoding_label.setStyleSheet(label_style)
        
        # 添加到状态栏
        self.status_bar.addWidget(self.char_count_label)
        self.status_bar.addWidget(self.selection_label)
        self.status_bar.addWidget(self.theme_label)
        self.status_bar.addPermanentWidget(self.encoding_label)

    def update_status_bar(self):
        text = self.editor.toPlainText()
        self.char_count_label.setText(f"字符: {len(text)}")
        self.selection_label.setText(f"选中: {len(self.editor.textCursor().selectedText())}")

        theme_info = PreviewThemes.get_theme_styles(self.preview_theme)
        self.theme_label.setText(f"预览主题: {theme_info['name']}")

    # ---------------- 编辑器样式 ----------------
    def update_editor_style(self):
        from qfluentwidgets import isDarkTheme
        is_dark = isDarkTheme()
        if is_dark:
            self.editor.setStyleSheet(f'''
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 8px 0px 0px 8px;
                padding: 10px;
                color: #ffffff;
                font-size: {self.font_size}px;
                selection-background-color: rgba(100, 149, 237, 0.3);
            ''')
            self.editor.setCursorWidth(3)
        else:
            self.editor.setStyleSheet(f'''
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 8px 0px 0px 8px;
                padding: 10px;
                color: #333333;
                font-size: {self.font_size}px;
                selection-background-color: rgba(100, 149, 237, 0.3);
            ''')
            self.editor.setCursorWidth(2)

    def _on_text_changed(self):
        if not self.has_file:
            return
        self.is_modified = True
        if self._preview_updating:
            self._preview_dirty = True
            return
        self._preview_timer.stop()
        self._preview_timer.start(self.PREVIEW_UPDATE_DELAY)

    def _do_preview_update(self):
        if not self.has_file:
            return
        if self._preview_updating:
            self._preview_dirty = True
            return
        self.update_preview()
        while self._preview_dirty:
            self._preview_dirty = False
            self.update_preview()

    # ---------------- 预览：Mica + 圆角 + 底部滚动修复 ----------------
    def update_preview(self):
        md_text = self.editor.toPlainText()
        current_md5 = QCryptographicHash.hash(md_text.encode('utf-8'), QCryptographicHash.Md5).toHex().data().decode()
        if current_md5 == self._last_md5 and self._cached_html_template:
            self._apply_cached_html()
            return
        self._preview_updating = True

        from qfluentwidgets import isDarkTheme
        is_dark = isDarkTheme()

        ts = PreviewThemes.get_theme_styles(self.preview_theme)
        bg = ts["background_color"]  # light=transparent -> 透出 Mica
        
        # 为默认主题添加系统主题适配
        if self.preview_theme == "light":
            if is_dark:
                # 深色系统主题下的默认主题配置
                text_color = "#e0e0e0"
                heading_color = "#ffffff"
                code_bg = "rgba(255, 255, 255, 0.1)"
                blockquote_bg = "rgba(255, 255, 255, 0.05)"
                scrollbar_track = "#3d3d3d"
                scrollbar_thumb = "#5d5d5d"
                scrollbar_thumb_hover = "#7d7d7d"
                link_color = "#64b5f6"
            else:
                # 浅色系统主题下的默认主题配置
                text_color = ts["text_color"]
                heading_color = ts["heading_color"]
                code_bg = ts["code_bg"]
                blockquote_bg = ts["blockquote_bg"]
                scrollbar_track = ts["scrollbar_track"]
                scrollbar_thumb = ts["scrollbar_thumb"]
                scrollbar_thumb_hover = ts["scrollbar_thumb_hover"]
                link_color = ts["link_color"]
        else:
            # 其他主题使用原有配置
            text_color = ts["text_color"]
            heading_color = ts["heading_color"]
            code_bg = ts["code_bg"]
            blockquote_bg = ts["blockquote_bg"]
            scrollbar_track = ts["scrollbar_track"]
            scrollbar_thumb = ts["scrollbar_thumb"]
            scrollbar_thumb_hover = ts["scrollbar_thumb_hover"]
            link_color = ts["link_color"]
        
        r = self.PREVIEW_RADIUS

        html = markdown.markdown(md_text, extensions=['fenced_code'])

        # 关键点：
        # - html/body 永远透明 -> 角落透出 Mica
        # - .content 用主题 bg（浅色 transparent 直接透出 Mica；其他主题为实色）
        # - .content 负责圆角裁剪（overflow:hidden）
        # - .scroll 负责滚动，避免底部显示不全
        # 使用字符串拼接构建HTML模板，避免f-string语法冲突
        styled_html = '''
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
  html, body {
    margin: 0;
    padding: 0;
    height: 100%;
    overflow: hidden;
    background: transparent !important;
    color: ''' + text_color + ''';
    font-family: Arial, sans-serif;
    font-size: ''' + str(self.font_size) + '''px;
  }

  .content {
    height: 100%;
    background: ''' + bg + ''';
    border-top-right-radius: ''' + str(r) + '''px;
    border-bottom-right-radius: ''' + str(r) + '''px;
    overflow: hidden;
  }

  .scroll {
    height: 100%;
    overflow-y: auto;
    box-sizing: border-box;
    padding: 20px 20px 36px 20px;
  }

  ::-webkit-scrollbar { width: 8px; height: 8px; }
  ::-webkit-scrollbar-track { background: ''' + scrollbar_track + '''; border-radius: 4px; }
  ::-webkit-scrollbar-thumb { background: ''' + scrollbar_thumb + '''; border-radius: 4px; }
  ::-webkit-scrollbar-thumb:hover { background: ''' + scrollbar_thumb_hover + '''; }

  h1,h2,h3,h4,h5,h6 { color: ''' + heading_color + '''; margin: 20px 0 10px; }
  p { margin: 0 0 10px; }

  code { background: ''' + code_bg + '''; padding: 2px 4px; border-radius: 3px; color: ''' + text_color + '''; }
  pre { position: relative; background: ''' + code_bg + '''; padding: 10px; border-radius: 5px; overflow-x: auto; margin: 10px 0; color: ''' + text_color + '''; }
  pre code { background: transparent; padding: 0; border-radius: 0; }
  .copy-button { position: absolute; top: 5px; right: 5px; background: rgba(255,255,255,0.2); border: 1px solid rgba(255,255,255,0.3); border-radius: 3px; padding: 2px 6px; font-size: 12px; cursor: pointer; color: ''' + text_color + '''; }
  .copy-button:hover { background: rgba(255,255,255,0.3); }
  .copy-button.copied { background: rgba(76,175,80,0.7); color: white; }

  blockquote {
    border-left: 4px solid rgba(100,149,237,0.5);
    margin: 10px 0;
    padding: 10px 15px;
    background: ''' + blockquote_bg + ''';
  }

  a { color: ''' + link_color + '''; text-decoration: none; }
  a:hover { text-decoration: underline; }

  table { border-collapse: collapse; width: 100%; margin: 10px 0; }
  th, td { border: 1px solid rgba(0,0,0,0.12); padding: 8px; text-align: left; }
  th { background: ''' + code_bg + '''; }
</style>
</head>
<body>
  <div class="content">
    <div class="scroll">
      ''' + html + '''
    </div>
  </div>
</body>
<script>
  // 为每个代码块添加复制按钮
  document.addEventListener("DOMContentLoaded", function() {
    var preElements = document.querySelectorAll("pre");
    preElements.forEach(function(pre) {
      // 创建复制按钮
      var copyButton = document.createElement("button");
      copyButton.className = "copy-button";
      copyButton.textContent = "复制";
      
      // 添加到代码块
      pre.appendChild(copyButton);
      
      // 复制功能
      copyButton.addEventListener("click", function() {
        var code = pre.querySelector("code");
        if (code) {
          var text = code.textContent;
          try {
            // 备用方法：创建临时文本域并复制（优先使用，避免剪贴板API的沙盒限制）
            var textArea = document.createElement("textarea");
            textArea.value = text;
            textArea.style.position = "fixed";
            textArea.style.left = "-999999px";
            textArea.style.top = "-999999px";
            document.body.appendChild(textArea);
            textArea.focus();
            textArea.select();
            var success = document.execCommand("copy");
            if (success) {
              copyButton.textContent = "已复制";
              copyButton.classList.add("copied");
              setTimeout(function() {
                copyButton.textContent = "复制";
                copyButton.classList.remove("copied");
              }, 2000);
            } else {
              throw new Error("复制失败");
            }
          } catch (err) {
            // 尝试使用剪贴板API作为备用
            if (navigator.clipboard && typeof navigator.clipboard.writeText === 'function') {
              navigator.clipboard.writeText(text).then(function() {
                copyButton.textContent = "已复制";
                copyButton.classList.add("copied");
                setTimeout(function() {
                  copyButton.textContent = "复制";
                  copyButton.classList.remove("copied");
                }, 2000);
              }).catch(function() {
                copyButton.textContent = "复制失败";
                setTimeout(function() {
                  copyButton.textContent = "复制";
                }, 2000);
              });
            } else {
              copyButton.textContent = "复制失败";
              setTimeout(function() {
                copyButton.textContent = "复制";
              }, 2000);
            }
          } finally {
            // 清理临时元素
            var textArea = document.querySelector("textarea[style*='-999999px']");
            if (textArea) {
              document.body.removeChild(textArea);
            }
          }
        }
      });
    });
  });
</script>
</html>
'''
        self.preview.setHtml(styled_html)
        self._cached_html_template = styled_html
        self._last_md5 = current_md5
        self._preview_updating = False
        if self._preview_dirty:
            self._preview_dirty = False
            QTimer.singleShot(0, self._do_preview_update)
        self._updatePreviewRoundMask()
        self.update_status_bar()

    def _apply_cached_html(self):
        if self._cached_html_template:
            self.preview.setHtml(self._cached_html_template)

    # ---------------- 基础功能 ----------------
    def open_file_dialog(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "Markdown Files (*.md);;All Files (*)")
        self.open_file(file_path)

    def save_file_dialog(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Markdown Files (*.md);;All Files (*)")
        self.save_file(file_path)

    def open_file(self, file_path):
        if file_path:
            self._preview_timer.stop()
            self._preview_updating = True
            self._last_md5 = None
            self._cached_html_template = None
            self.current_file_path = file_path
            self.is_modified = False
            self.has_file = True
            self._add_to_recent_files(file_path)
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            self._hide_welcome_page()
            self.editor.setPlainText(content)
            self._start_auto_save_timer()
            self._update_command_bar_enabled()
            QTimer.singleShot(0, self._do_preview_update)

    def save_file(self, file_path):
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(self.editor.toPlainText())
            self.current_file_path = file_path
            self.is_modified = False
            self._start_auto_save_timer()

    def new_file(self):
        self._preview_timer.stop()
        self._preview_updating = True
        self._last_md5 = None
        self._cached_html_template = None
        self.current_file_path = None
        self.is_modified = False
        self.has_file = True
        self._hide_welcome_page()
        self.editor.clear()
        self._start_auto_save_timer()
        self._update_command_bar_enabled()
        QTimer.singleShot(0, self._do_preview_update)

    def copy(self):
        self.editor.copy()

    def paste(self):
        from PyQt5.QtWidgets import QApplication
        import tempfile, uuid

        clipboard = QApplication.clipboard()
        if clipboard.mimeData().hasImage():
            image = clipboard.image()
            if not image.isNull():
                temp_dir = tempfile.gettempdir()
                file_name = f"image_{uuid.uuid4().hex}.png"
                file_path = os.path.join(temp_dir, file_name)
                if image.save(file_path, "PNG"):
                    file_url = self._local_path_to_url(file_path)
                    self.editor.textCursor().insertText(f"![{file_name}]({file_url})")
                    self.update_preview()
                    return
        self.editor.paste()

    def insert_image(self):
        if not self.has_file:
            return
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Image", "",
            "Image Files (*.png *.jpg *.jpeg *.gif *.bmp *.svg);;All Files (*)"
        )
        if file_path:
            image_name = os.path.basename(file_path)
            file_url = self._local_path_to_url(file_path)
            self.editor.textCursor().insertText(f"![{image_name}]({file_url})")
            self.update_preview()

    def _local_path_to_url(self, path):
        import urllib.parse
        return urllib.parse.quote(path.replace('\\', '/'), safe=':/')

    def toggle_fullscreen(self):
        if not self.has_file:
            return
        if not self.is_fullscreen:
            self.editor.hide()
            self.splitter.setSizes([0, self.splitter.width()])
            self.editor_layout.setContentsMargins(0, 0, 0, 0)
            self.card_container_layout.setContentsMargins(0, 0, 0, 0)
            self.is_fullscreen = True
            self.is_editor_fullscreen = False
        else:
            self.editor.show()
            w = self.splitter.width()
            self.splitter.setSizes([w // 2, w // 2])
            self.editor_layout.setContentsMargins(1, 1, 1, 1)
            self.card_container_layout.setContentsMargins(5, 5, 5, 5)
            self.is_fullscreen = False

        self._updatePreviewRoundMask()

    def toggle_editor_fullscreen(self):
        if not self.has_file:
            return
        if not self.is_editor_fullscreen:
            self.preview_container.hide()
            w = self.splitter.width()
            self.splitter.setSizes([w, 0])
            self.editor_layout.setContentsMargins(0, 0, 0, 0)
            self.card_container_layout.setContentsMargins(0, 0, 0, 0)
            self.scroll_area.setStyleSheet("QScrollArea{background:transparent;border:none;}")
            self.is_editor_fullscreen = True
            self.is_fullscreen = False
        else:
            self.preview_container.show()
            w = self.splitter.width()
            self.splitter.setSizes([w // 2, w // 2])
            self.editor_layout.setContentsMargins(1, 1, 1, 1)
            self.card_container_layout.setContentsMargins(5, 5, 5, 5)
            self.scroll_area.setStyleSheet("")
            self.is_editor_fullscreen = False

        self._updatePreviewRoundMask()
    
    def zoom_in(self):
        """放大字体"""
        if not self.has_file:
            return
        self.font_size = min(self.font_size + 2, 32)  # 最大32px
        self.update_editor_style()
        self.update_preview()

    def zoom_out(self):
        """缩小字体"""
        if not self.has_file:
            return
        self.font_size = max(self.font_size - 2, 8)   # 最小8px
        self.update_editor_style()
        self.update_preview()

    def zoom_reset(self):
        """重置字体大小"""
        if not self.has_file:
            return
        self.font_size = 16  # 恢复默认字体大小
        self.update_editor_style()
        self.update_preview()
    
    def eventFilter(self, obj, event):
        """事件过滤器，用于捕获编辑器的鼠标滚轮事件"""
        from PyQt5.QtCore import Qt, QEvent
        if obj == self.editor and event.type() == QEvent.Wheel:
            if event.modifiers() & Qt.ControlModifier:
                # Ctrl+滚轮调节字体大小
                delta = event.angleDelta().y()
                if delta > 0:
                    # 放大字体
                    self.font_size = min(self.font_size + 2, 32)  # 最大32px
                else:
                    # 缩小字体
                    self.font_size = max(self.font_size - 2, 8)   # 最小8px
                
                # 更新编辑器和预览的字体大小
                self.update_editor_style()
                self.update_preview()
                return True  # 事件已处理
        return False  # 事件未处理，传递给其他处理者

    # ---------------- 导出（保持原逻辑，略） ----------------
    def export_file(self):
        if not self.has_file:
            return
        file_path, file_type = QFileDialog.getSaveFileName(
            self, "Export File", "",
            "PDF Files (*.pdf);;Word Files (*.docx);;HTML Files (*.html);;All Files (*)"
        )
        if not file_path:
            return

        markdown_text = self.editor.toPlainText()

        if not file_path.endswith(('.pdf', '.docx', '.html')):
            if 'PDF Files' in file_type:
                file_path += '.pdf'
            elif 'Word Files' in file_type:
                file_path += '.docx'
            elif 'HTML Files' in file_type:
                file_path += '.html'

        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            self.export_to_pdf(file_path, markdown_text)
        elif ext == '.docx':
            self.export_to_word(file_path, markdown_text)
        elif ext == '.html':
            self.export_to_html(file_path, markdown_text)

    def export_to_pdf(self, file_path, markdown_text):
        if not HAS_EXPORT_LIBS:
            self._show_info_dialog("导出错误", "fpdf 库未安装，请运行: pip install fpdf")
            return
        try:
            import re
            import warnings
            warnings.filterwarnings('ignore')

            pdf = FPDF()
            pdf.add_page()

            font_path = 'C:/Windows/Fonts/simhei.ttf'
            code_font_path = 'C:/Windows/Fonts/consola.ttf'

            if os.path.exists(font_path):
                pdf.add_font('SimHei', '', font_path, uni=True)
                pdf.add_font('SimHei', 'B', font_path, uni=True)
                font_regular = 'SimHei'
                font_bold = 'SimHei'
            else:
                font_regular = 'Arial'
                font_bold = 'Arial'

            if os.path.exists(code_font_path):
                pdf.add_font('Consolas', '', code_font_path, uni=True)
                code_font = 'Consolas'
            else:
                code_font = font_regular

            def contains_chinese(text):
                return bool(re.search(r'[\u4e00-\u9fff]', text))

            def get_code_font(text):
                if contains_chinese(text) and os.path.exists(font_path):
                    return 'SimHei'
                return code_font

            lines = markdown_text.split('\n')
            i = 0
            in_code_block = False
            while i < len(lines):
                line = lines[i]

                if line.strip().startswith('```'):
                    if not in_code_block:
                        in_code_block = True
                        i += 1
                        continue
                    else:
                        in_code_block = False
                        i += 1
                        continue

                if in_code_block:
                    current_font = get_code_font(line)
                    pdf.set_font(current_font, size=11)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.multi_cell(0, 6, line, fill=True)
                    i += 1
                    continue

                if line.startswith('# '):
                    pdf.set_font(font_bold, size=20)
                    pdf.multi_cell(0, 10, line[2:])
                elif line.startswith('## '):
                    pdf.set_font(font_bold, size=16)
                    pdf.multi_cell(0, 8, line[3:])
                elif line.startswith('### '):
                    pdf.set_font(font_bold, size=14)
                    pdf.multi_cell(0, 7, line[4:])
                elif line.startswith('- ') or line.startswith('* '):
                    pdf.set_font(font_regular, size=12)
                    pdf.multi_cell(0, 6, f"  •  {line[2:]}")
                elif re.match(r'^\d+\.\s', line):
                    match = re.match(r'^(\d+)\.\s(.*)', line)
                    if match:
                        pdf.set_font(font_regular, size=12)
                        pdf.multi_cell(0, 6, f"  {match.group(1)}. {match.group(2)}")
                    else:
                        pdf.set_font(font_regular, size=12)
                        pdf.multi_cell(0, 6, line)
                elif line.startswith('>'):
                    pdf.set_font(font_regular, size=11)
                    pdf.set_fill_color(245, 245, 245)
                    pdf.multi_cell(0, 6, line[1:], fill=True)
                elif line.strip() == '':
                    pdf.ln(3)
                else:
                    pdf.set_font(font_regular, size=12)
                    pdf.multi_cell(0, 6, line)
                i += 1

            pdf_content = pdf.output(dest='S')
            with open(file_path, 'wb') as f:
                f.write(pdf_content.encode('latin-1'))
            self._show_info_dialog("导出成功", f"PDF 已成功导出到:\n{file_path}")
        except Exception as e:
            error_msg = f"导出 PDF 时出错:\n{str(e)}"
            print(f"Error exporting to PDF: {e}")
            import traceback
            traceback.print_exc()
            self._show_info_dialog("导出错误", error_msg)

    def _parse_inline_style(self, text, pdf, font_regular, font_bold, code_font):
        import re
        result = []
        pattern = re.compile(r'\*\*(.+?)\*\*(.+?)', re.DOTALL)
        pattern2 = re.compile(r'\*(.+?)\*(.+?)', re.DOTALL)
        pattern3 = re.compile(r'`(.+?)`(.+?)', re.DOTALL)

        remaining = text
        while remaining:
            bold_match = pattern.match(remaining)
            italic_match = pattern2.match(remaining)
            code_match = pattern3.match(remaining)

            if bold_match:
                result.append(('bold', bold_match.group(1)))
                remaining = bold_match.group(2)
            elif italic_match:
                result.append(('italic', italic_match.group(1)))
                remaining = italic_match.group(2)
            elif code_match:
                result.append(('code', code_match.group(1)))
                remaining = code_match.group(2)
            else:
                if remaining:
                    result.append(('normal', remaining))
                break

        return ''.join([item[1] for item in result])

    def export_to_word(self, file_path, markdown_text):
        if not HAS_EXPORT_LIBS:
            print("Error: python-docx not installed. pip install python-docx")
            return
        try:
            doc = Document()
            for line in markdown_text.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    if line.strip():
                        doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            print(f"Error exporting to Word: {e}")

    def export_to_html(self, file_path, markdown_text):
        try:
            html = markdown.markdown(markdown_text, extensions=['fenced_code'])
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"<!doctype html><meta charset='utf-8'><body>{html}</body>")
        except Exception as e:
            print(f"Error exporting to HTML: {e}")
