import os
import markdown


class ExportController:
    HAS_EXPORT_LIBS = False

    try:
        from fpdf import FPDF
        from docx import Document
        HAS_EXPORT_LIBS = True
    except ImportError:
        pass

    @staticmethod
    def export_pdf(file_path, markdown_text):
        if not ExportController.HAS_EXPORT_LIBS:
            return False, "fpdf 库未安装，请运行: pip install fpdf"
        
        try:
            import re
            import warnings
            from fpdf import FPDF
            warnings.filterwarnings('ignore')

            pdf = FPDF()
            pdf.add_page()

            font_path = 'C:/Windows/Fonts/simhei.ttf'
            code_font_path = 'C:/Windows/Fonts/consola.ttf'

            font_regular = 'Arial'
            font_bold = 'Arial'
            code_font = 'Arial'

            if os.path.exists(font_path):
                pdf.add_font('SimHei', '', font_path, uni=True)
                pdf.add_font('SimHei', 'B', font_path, uni=True)
                font_regular = 'SimHei'
                font_bold = 'SimHei'

            if os.path.exists(code_font_path):
                pdf.add_font('Consolas', '', code_font_path, uni=True)
                code_font = 'Consolas'

            def contains_chinese(text):
                return bool(re.search(r'[\u4e00-\u9fff]', text))

            def get_code_font(text):
                if contains_chinese(text) and os.path.exists(font_path):
                    return 'SimHei'
                return code_font

            lines = markdown_text.split('\n')
            i = 0
            in_code_block = False
            while i < len(lines):
                line = lines[i]

                if line.strip().startswith('```'):
                    in_code_block = not in_code_block
                    i += 1
                    continue

                if in_code_block:
                    current_font = get_code_font(line)
                    pdf.set_font(current_font, size=11)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.multi_cell(0, 6, line, fill=True)
                elif line.startswith('# '):
                    pdf.set_font(font_bold, size=20)
                    pdf.multi_cell(0, 10, line[2:])
                elif line.startswith('## '):
                    pdf.set_font(font_bold, size=16)
                    pdf.multi_cell(0, 8, line[3:])
                elif line.startswith('### '):
                    pdf.set_font(font_bold, size=14)
                    pdf.multi_cell(0, 7, line[4:])
                elif line.startswith('- ') or line.startswith('* '):
                    pdf.set_font(font_regular, size=12)
                    pdf.multi_cell(0, 6, f"  •  {line[2:]}")
                elif re.match(r'^\d+\.\s', line):
                    match = re.match(r'^(\d+)\.\s(.*)', line)
                    if match:
                        pdf.set_font(font_regular, size=12)
                        pdf.multi_cell(0, 6, f"  {match.group(1)}. {match.group(2)}")
                    else:
                        pdf.set_font(font_regular, size=12)
                        pdf.multi_cell(0, 6, line)
                elif line.startswith('>'):
                    pdf.set_font(font_regular, size=11)
                    pdf.set_fill_color(245, 245, 245)
                    pdf.multi_cell(0, 6, line[1:], fill=True)
                elif line.strip() == '':
                    pdf.ln(3)
                else:
                    pdf.set_font(font_regular, size=12)
                    pdf.multi_cell(0, 6, line)
                i += 1

            pdf_content = pdf.output(dest='S')
            with open(file_path, 'wb') as f:
                f.write(pdf_content.encode('latin-1'))
            return True, f"PDF 已成功导出到:\n{file_path}"
        except Exception as e:
            return False, f"导出 PDF 时出错:\n{str(e)}"

    @staticmethod
    def export_word(file_path, markdown_text):
        if not ExportController.HAS_EXPORT_LIBS:
            return False, "python-docx 库未安装，请运行: pip install python-docx"
        
        try:
            from docx import Document
            doc = Document()
            for line in markdown_text.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    if line.strip():
                        doc.add_paragraph(line)
            doc.save(file_path)
            return True, f"Word 文档已成功导出到:\n{file_path}"
        except Exception as e:
            return False, f"导出 Word 文档时出错:\n{str(e)}"

    @staticmethod
    def export_html(file_path, markdown_text):
        try:
            html = markdown.markdown(markdown_text, extensions=['fenced_code'])
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"<!doctype html><meta charset='utf-8'><body>{html}</body>")
            return True, f"HTML 文件已成功导出到:\n{file_path}"
        except Exception as e:
            return False, f"导出 HTML 文件时出错:\n{str(e)}"
